from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from pypdf import PdfReader


class ResumeParseError(RuntimeError):
    pass


def parse_resume_bytes(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    try:
        if suffix == ".pdf":
            reader = PdfReader(io.BytesIO(data))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        elif suffix == ".docx":
            doc = Document(io.BytesIO(data))
            blocks = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    blocks.append(" | ".join(cell.text for cell in row.cells))
            text = "\n".join(blocks)
        elif suffix in {".txt", ".md"}:
            text = data.decode("utf-8", errors="ignore")
        else:
            raise ResumeParseError("仅支持 PDF、DOCX、TXT 和 Markdown。")
    except Exception as exc:
        if isinstance(exc, ResumeParseError):
            raise
        raise ResumeParseError(f"简历解析失败：{exc}") from exc

    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if len(text) < 80:
        raise ResumeParseError("提取到的文字太少。若 PDF 是扫描件，请先转成可复制文字的 PDF。")
    return text
